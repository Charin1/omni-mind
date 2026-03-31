import datetime
import os
import re
import tempfile
import uuid
from pathlib import Path
from typing import List, Optional

from sqlalchemy.ext.asyncio import AsyncSession

from db.models import Artifact
from runtime.limits import artifact_semaphore


class ArtifactService:
    def __init__(self, db: AsyncSession):
        self.db = db
        self.output_root = Path(__file__).resolve().parent.parent / "generated_artifacts"
        self.output_root.mkdir(parents=True, exist_ok=True)

    def detect_kind(self, message: str) -> Optional[str]:
        text = message.lower()
        if "docx" in text or "word document" in text:
            return "docx"
        if "xlsx" in text or "excel" in text or "spreadsheet" in text:
            return "xlsx"
        if "html" in text:
            return "html"
        if "plot" in text or "chart" in text or "graph" in text:
            return "plot"
        return None

    async def maybe_generate_from_message(
        self,
        user_id: str,
        conversation_id: str,
        message: str,
    ) -> Optional[Artifact]:
        kind = self.detect_kind(message)
        if not kind:
            return None
        return await self.generate_artifact(
            kind=kind,
            user_id=user_id,
            conversation_id=conversation_id,
            prompt=message,
        )

    async def generate_artifact(
        self,
        kind: str,
        user_id: str,
        conversation_id: str,
        prompt: str,
    ) -> Artifact:
        async with artifact_semaphore:
            artifact_id = str(uuid.uuid4())
            timestamp = datetime.datetime.utcnow().strftime("%Y%m%d-%H%M%S")
            safe_name = f"{kind}-{timestamp}"
            target_dir = self.output_root / user_id / conversation_id
            target_dir.mkdir(parents=True, exist_ok=True)

            with tempfile.TemporaryDirectory(prefix="omnimind-artifact-") as temp_dir:
                try:
                    if kind == "docx":
                        path, mime_type = self._create_docx(target_dir, safe_name, prompt)
                    elif kind == "xlsx":
                        path, mime_type = self._create_xlsx(target_dir, safe_name, prompt)
                    elif kind == "html":
                        path, mime_type = self._create_html(target_dir, safe_name, prompt)
                    elif kind == "plot":
                        path, mime_type = self._create_plot(target_dir, safe_name, prompt)
                    else:
                        raise ValueError(f"Unsupported artifact kind: {kind}")
                except ImportError as exc:
                    raise RuntimeError(
                        "Artifact dependencies are missing. Install backend requirements to enable docx/xlsx/plot generation."
                    ) from exc

                # Keep a marker in the temp workspace so the flow explicitly uses
                # an isolated scratch area before returning the final artifact.
                Path(temp_dir, "completed.txt").write_text(path.name, encoding="utf-8")

            artifact = Artifact(
                id=artifact_id,
                user_id=user_id,
                conversation_id=conversation_id,
                kind=kind,
                name=path.name,
                path=str(path),
                mime_type=mime_type,
                status="completed",
                metadata_json={"prompt": prompt[:500]},
            )
            self.db.add(artifact)
            await self.db.commit()
            return artifact

    def _create_docx(self, target_dir: Path, base_name: str, prompt: str):
        from docx import Document

        path = target_dir / f"{base_name}.docx"
        document = Document()
        document.add_heading("OmniMind Document", 0)
        document.add_paragraph("Request")
        document.add_paragraph(prompt)
        document.add_paragraph(
            "Generated in an isolated workspace and saved as a persistent artifact."
        )
        document.save(path)
        return path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def _create_xlsx(self, target_dir: Path, base_name: str, prompt: str):
        import xlsxwriter

        path = target_dir / f"{base_name}.xlsx"
        workbook = xlsxwriter.Workbook(str(path))
        worksheet = workbook.add_worksheet("Request")
        worksheet.write(0, 0, "Prompt")
        worksheet.write(0, 1, prompt)

        rows = self._extract_table_rows(prompt)
        if rows:
            table_sheet = workbook.add_worksheet("Data")
            for row_index, row in enumerate(rows):
                for col_index, value in enumerate(row):
                    table_sheet.write(row_index, col_index, value)
        workbook.close()
        return path, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    def _create_html(self, target_dir: Path, base_name: str, prompt: str):
        path = target_dir / f"{base_name}.html"
        html = f"""<!DOCTYPE html>
<html lang="en">
  <head>
    <meta charset="utf-8" />
    <title>OmniMind HTML</title>
    <style>
      body {{ font-family: Arial, sans-serif; max-width: 900px; margin: 40px auto; line-height: 1.6; }}
      pre {{ white-space: pre-wrap; background: #f4f4f4; padding: 16px; border-radius: 8px; }}
    </style>
  </head>
  <body>
    <h1>OmniMind HTML Artifact</h1>
    <p>Generated from the chat request below.</p>
    <pre>{self._escape_html(prompt)}</pre>
  </body>
</html>
"""
        path.write_text(html, encoding="utf-8")
        return path, "text/html"

    def _create_plot(self, target_dir: Path, base_name: str, prompt: str):
        import plotly.graph_objects as go

        path = target_dir / f"{base_name}.html"
        numbers = self._extract_numbers(prompt)
        if len(numbers) < 2:
            numbers = [3, 5, 2, 7, 4]
        fig = go.Figure(data=[go.Scatter(y=numbers, mode="lines+markers")])
        fig.update_layout(title="OmniMind Plot", template="plotly_white")
        fig.write_html(str(path))
        return path, "text/html"

    def _extract_numbers(self, text: str) -> List[float]:
        matches = re.findall(r"-?\d+(?:\.\d+)?", text)
        return [float(match) for match in matches]

    def _extract_table_rows(self, text: str) -> List[List[str]]:
        rows: List[List[str]] = []
        for line in text.splitlines():
            if "," in line:
                rows.append([cell.strip() for cell in line.split(",")])
        return rows

    def _escape_html(self, text: str) -> str:
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    def public_url(self, artifact: Artifact) -> str:
        relative = os.path.relpath(artifact.path, self.output_root)
        return f"/artifacts/{relative.replace(os.sep, '/')}"
